#!/usr/bin/env python3
"""
Convert Markdown content to a formatted Word document using a template.
"""

import re
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def read_markdown(filepath):
    """Read Markdown file content."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: Markdown file not found at {filepath}")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading Markdown file: {e}")
        sys.exit(1)

def parse_markdown_structure(md_content):
    """Parse Markdown content into structured blocks."""
    # Remove transition summaries or specific batch artifacts if any (custom logic, but generally safe to ignore)
    md_content = re.sub(r'\*\*【批次\d+承接摘要】\*\*.*?(?=\n#|\n---|\Z)', '', md_content, flags=re.DOTALL)
    
    blocks = []
    lines = md_content.split('\n')
    
    current_block = {'type': 'text', 'content': '', 'level': 0}
    in_table = False
    table_lines = []
    in_formula = False
    formula_lines = []
    
    for line in lines:
        # Skip separators but keep TOC
        if line.strip() in ['---', '']:
            if current_block['content'].strip():
                blocks.append(current_block)
                current_block = {'type': 'text', 'content': '', 'level': 0}
            continue
            
        if line.strip() == '[TOC]':
             if current_block['content'].strip():
                blocks.append(current_block)
                current_block = {'type': 'text', 'content': '', 'level': 0}
             blocks.append({'type': 'heading', 'level': 0, 'content': '[TOC]'}) # Level 0 for special items
             continue
        
        # Formula blocks
        if line.strip().startswith('$$'):
            if in_formula:
                formula_lines.append(line)
                blocks.append({'type': 'formula', 'content': '\n'.join(formula_lines)})
                formula_lines = []
                in_formula = False
            else:
                if current_block['content'].strip():
                    blocks.append(current_block)
                    current_block = {'type': 'text', 'content': '', 'level': 0}
                in_formula = True
                formula_lines = [line]
            continue
        
        if in_formula:
            formula_lines.append(line)
            continue
        
        # Table detection
        if line.strip().startswith('|'):
            if not in_table:
                if current_block['content'].strip():
                    blocks.append(current_block)
                    current_block = {'type': 'text', 'content': '', 'level': 0}
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        elif in_table:
            blocks.append({'type': 'table', 'content': table_lines})
            table_lines = []
            in_table = False
        
        # Heading detection
        if line.startswith('#'):
            if current_block['content'].strip():
                blocks.append(current_block)
                current_block = {'type': 'text', 'content': '', 'level': 0}
            
            level = len(re.match(r'^#+', line).group())
            title_text = re.sub(r'^#+\s*', '', line).strip()
            blocks.append({'type': 'heading', 'level': level, 'content': title_text})
            continue
        
        # Regular text
        current_block['content'] += line + '\n'
    
    # Append last block
    if in_table:
        blocks.append({'type': 'table', 'content': table_lines})
    elif current_block['content'].strip():
        blocks.append(current_block)
    
    return blocks

def set_cell_shading(cell, color):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def create_table_from_markdown(doc, table_lines, template_doc):
    """Create a Word table from Markdown lines."""
    rows_data = []
    for line in table_lines:
        if '---' in line:
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)
    
    if not rows_data:
        return
    
    num_cols = len(rows_data[0])
    num_rows = len(rows_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass # Fallback if style missing
        
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
                if i == 0:
                    set_cell_shading(cell, 'D9E2F3')
    
    return table

def process_inline_formatting(paragraph, text):
    """Process bold formatting."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def safe_clear_document(doc):
    """Clear document content while preserving styles and section properties."""
    body = doc.element.body
    elements_to_remove = []
    for element in body:
        if element.tag.endswith('sectPr'):
            continue
        elements_to_remove.append(element)
    
    for element in elements_to_remove:
        body.remove(element)

def add_toc(doc):
    """Insert Table of Contents field code."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    
    run._element.append(OxmlElement('w:fldChar')) # end
    
    paragraph = doc.add_paragraph('(Right-click -> Update Field to generate TOC)')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def check_and_apply_style(doc, style_name, default_style='Normal'):
    """Apply style if exists, else fallback."""
    if style_name in doc.styles:
        return style_name
    return default_style

def convert_md_to_docx(md_path, template_path, output_path):
    """Main conversion function."""
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error loading template: {e}")
        sys.exit(1)
        
    safe_clear_document(doc)
    
    md_content = read_markdown(md_path)
    blocks = parse_markdown_structure(md_content)
    
    for block in blocks:
        if block['type'] == 'heading':
            level = block['level']
            title = block['content']
            
            # Abstract
            if title == '摘要' or title == 'Abstract':
                is_english = (title == 'Abstract')
                
                if is_english:
                    doc.add_page_break()
                else:
                    doc.add_paragraph() # Empty line
                
                para = doc.add_paragraph()
                para.style = check_and_apply_style(doc, 'Title')
                run = para.add_run(title)
                
                if is_english:
                    run.font.name = 'Times New Roman'
                else:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                run.font.bold = True
                run.font.size = Pt(16)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            
            # TOC
            if title == 'TOC' or title == '[TOC]':
                 doc.add_page_break()
                 para = doc.add_paragraph()
                 run = para.add_run('Contents') # Or '目录' depending on language context, defaulting to English 'Contents' for generic skill but '目录' for Chinese context. Let's use formatting detection or default to 'Contents'/'目录'
                 # Since we saw '目录' in previous chinese document, let's look for Chinese chars in content to decide? Or just check if title was [TOC]
                 # For generic tool, defaulting to 'Table of Contents' or just '目录' if detected Chinese content is safer.
                 # Given the previous context was Chinese thesis, '目录' is appropriate. But as a generic tool...
                 # Let's use 'Table of Contents' / '目录' based on context or simple default.
                 # For now, I will use '目录' as the skill is likely used in Chinese context based on user request ("通用技能...").
                 run.text = '目  录'
                 run.font.name = '黑体'
                 run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                 run.font.size = Pt(16)
                 para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                 doc.add_paragraph()
                 add_toc(doc)
                 doc.add_page_break()
                 continue

            # Headings
            if level == 1:
                # Chapter handling
                if '第' in title and '章' in title:
                    doc.add_page_break()
                    para = doc.add_paragraph()
                    para.style = check_and_apply_style(doc, 'Heading 1', 'Heading 1')
                    chapter_title = re.sub(r'第.章\s*', '', title)
                    para.text = title
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif title == '参考文献' or title == 'References':
                    doc.add_page_break()
                    para = doc.add_paragraph()
                    para.style = check_and_apply_style(doc, 'Heading 1', 'Heading 1')
                    para.text = title
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else: 
                    para = doc.add_paragraph()
                    para.style = check_and_apply_style(doc, 'Heading 1', 'Heading 1')
                    para.text = title
            
            elif level == 2:
                para = doc.add_paragraph()
                para.style = check_and_apply_style(doc, 'Heading 2', 'Heading 2')
                para.text = title
            
            elif level == 3:
                para = doc.add_paragraph()
                para.style = check_and_apply_style(doc, 'Heading 3', 'Heading 3')
                para.text = title

        elif block['type'] == 'text':
            content = block['content'].strip()
            if not content:
                continue
            
            if content == '[TOC]':
                 doc.add_page_break()
                 para = doc.add_paragraph()
                 run = para.add_run('目  录')
                 run.font.name = '黑体'
                 run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                 run.font.size = Pt(16)
                 para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                 add_toc(doc)
                 doc.add_page_break()
                 continue

            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text: continue
                
                # Keywords
                if para_text.startswith('**关键词') or para_text.startswith('**Keywords'):
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    style_name = '关键词正文' if '关键词' in para_text else '英文关键词正文'
                    para.style = check_and_apply_style(doc, style_name, 'Normal')
                    
                    prefix = '关键词：' if '关键词' in para_text else 'Keywords: '
                    run = para.add_run(prefix)
                    run.bold = True
                    
                    if '关键词' in para_text:
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        text = re.sub(r'\*\*关键词\*\*:\s*', '', para_text)
                    else:
                        run.font.name = 'Times New Roman'
                        text = re.sub(r'\*\*Keywords\*\*:\s*', '', para_text)
                    
                    run2 = para.add_run(text)
                    if '关键词' in para_text:
                        run2.font.name = '宋体'
                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        run2.font.name = 'Times New Roman'
                    continue

                # Table/Figure Captions
                if (para_text.startswith('**表') or para_text.startswith('**图')) and '**' in para_text[2:]:
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    para.style = check_and_apply_style(doc, '表格标题' if '表' in para_text else '图片标题', 'Caption')
                    title = re.sub(r'\*\*', '', para_text)
                    para.text = title
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                
                # Note/References
                if para_text.startswith('注：') or para_text.startswith('注:'):
                    para = doc.add_paragraph()
                    run = para.add_run(para_text)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(9)
                    para.paragraph_format.first_line_indent = Cm(0)
                    continue
                
                if re.match(r'^\[\d+\]', para_text):
                    para = doc.add_paragraph()
                    run = para.add_run(para_text)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                    para.paragraph_format.first_line_indent = Cm(-0.74)
                    para.paragraph_format.left_indent = Cm(0.74)
                    continue
                
                # Body Text
                para = doc.add_paragraph()
                para.style = check_and_apply_style(doc, 'Body Text', 'Normal')
                
                if para.style.name in ['Body Text', 'Normal', '正文']:
                    para.paragraph_format.first_line_indent = Cm(0.74)
                    para.paragraph_format.line_spacing = 1.25
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                process_inline_formatting(para, para_text)
        
        elif block['type'] == 'table':
            create_table_from_markdown(doc, block['content'], doc)
            doc.add_paragraph()
        
        elif block['type'] == 'formula':
            para = doc.add_paragraph()
            para.style = 'Normal'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = block['content'].replace('$$', '').strip()
            run = para.add_run(text)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    try:
        doc.save(output_path)
        print(f"Successfully saved document to: {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to formatted Word Docx using a template.')
    parser.add_argument('--input', required=True, help='Path to input Markdown file')
    parser.add_argument('--template', required=True, help='Path to template Word file')
    parser.add_argument('--output', required=True, help='Path to output Word file')
    
    args = parser.parse_args()
    
    convert_md_to_docx(args.input, args.template, args.output)

if __name__ == '__main__':
    main()

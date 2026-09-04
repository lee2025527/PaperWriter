#!/usr/bin/env python3
"""
Reorder and renumber references by first appearance in the body (GB/T 7714 顺序编码制).

- Body citations and reference list are renumbered 1, 2, 3, ... by first occurrence.
- Reference list paragraph order is reordered to match.
- Supports multiple list-entry prefix styles: [n], [n]., n., n  (preserved when updating).
"""

import re
from typing import Dict, Iterable, List, Optional, Set, Tuple

from docx import Document
from docx.oxml.ns import qn

# In-text citation: [n]
REF_PATTERN = re.compile(r"\[(\d{1,4})\]")

# Reference list line: various prefixes (capture number; format fn gives new prefix for a number)
REF_LINE_PATTERNS = [
    (re.compile(r"^\[(\d{1,4})\]\.?\s*(.*)$"), lambda n: f"[{n}] "),   # [1] or [1].
    (re.compile(r"^(\d{1,4})\.\s*(.*)$"), lambda n: f"{n}. "),         # 1.
    (re.compile(r"^(\d{1,4})\s+(.*)$"), lambda n: f"{n} "),            # 1 
]


def _match_ref_line(text: str) -> Optional[Tuple[str, str, str, int]]:
    """Return (num, prefix_for_num, rest, pattern_index) or None. pattern_index selects REF_LINE_PATTERNS."""
    text = text.strip()
    for idx, (pat, fmt_fn) in enumerate(REF_LINE_PATTERNS):
        m = pat.match(text)
        if m:
            num = m.group(1)
            rest = m.group(2) if len(m.groups()) > 1 and m.group(2) is not None else ""
            return (num, fmt_fn(num), rest, idx)
    return None


def _find_reference_section_start(doc: Document, ref_heading: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() != ref_heading:
            continue
        for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
            nxt = (doc.paragraphs[j].text or "").strip()
            if _match_ref_line(nxt) is not None:
                return i
    return -1


def _find_reference_section_end(doc: Document, start_idx: int, stop_headings: Set[str]) -> int:
    if start_idx < 0:
        return -1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = (doc.paragraphs[i].text or "").strip()
        if text in stop_headings:
            return i
    return len(doc.paragraphs)


def _iter_reference_paragraphs(doc: Document, start_idx: int, end_idx: int) -> Iterable:
    for i in range(start_idx + 1, end_idx):
        yield doc.paragraphs[i]


def _find_reference_heading_by_index(doc: Document, idx: int):
    if 0 <= idx < len(doc.paragraphs):
        return doc.paragraphs[idx]
    return None


def _reference_entries_in_list(doc: Document, start_idx: int, end_idx: int) -> List[Tuple[str, str, str, int]]:
    """Return list of (num, prefix, rest, pattern_index) for each ref line."""
    result: List[Tuple[str, str, str, int]] = []
    for p in _iter_reference_paragraphs(doc, start_idx, end_idx):
        m = _match_ref_line((p.text or "").strip())
        if m:
            result.append(m)
    return result


def _iter_run_elements(p) -> Iterable:
    for child in p._p.iterchildren():
        if child.tag == qn("w:r"):
            yield child
        elif child.tag == qn("w:hyperlink"):
            for r in child.iterchildren():
                if r.tag == qn("w:r"):
                    yield r


def _extract_ref_numbers_in_order(p, ref_numbers: Set[str]) -> List[str]:
    numbers: List[str] = []
    in_field = False
    field_num: Optional[str] = None
    for r in _iter_run_elements(p):
        fld_char = r.find(qn("w:fldChar"))
        instr = r.find(qn("w:instrText"))
        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_field = True
                field_num = None
            elif fld_type == "end":
                if field_num:
                    numbers.append(field_num)
                in_field = False
                field_num = None
        if in_field and instr is not None and instr.text:
            match = re.search(r"REF\\s+([A-Za-z0-9_]+)", instr.text)
            if match:
                name = match.group(1)
                if name.startswith("ref_"):
                    num = name.replace("ref_", "", 1)
                    if num in ref_numbers:
                        field_num = num
    if numbers:
        return numbers
    text = p.text or ""
    for m in REF_PATTERN.finditer(text):
        if m.group(1) in ref_numbers:
            numbers.append(m.group(1))
    return numbers


def _build_first_appearance_order(doc: Document, stop_idx: int, ref_numbers: Set[str]) -> List[str]:
    order: List[str] = []
    seen: Set[str] = set()
    for i in range(0, stop_idx):
        p = doc.paragraphs[i]
        nums = _extract_ref_numbers_in_order(p, ref_numbers)
        for num in nums:
            if num not in seen:
                order.append(num)
                seen.add(num)
    return order


def _update_bookmarked_number(p, new_num: str) -> bool:
    in_bookmark = False
    updated = False
    for child in p._p.iterchildren():
        if child.tag == qn("w:bookmarkStart"):
            name = child.get(qn("w:name")) or ""
            if name.startswith("ref_"):
                in_bookmark = True
        elif child.tag == qn("w:bookmarkEnd") and in_bookmark:
            in_bookmark = False
        elif in_bookmark and child.tag == qn("w:r"):
            for t in child.findall(qn("w:t")):
                t.text = new_num
                updated = True
    return updated


def _update_reference_number_text(p, old_num: str, new_num: str, new_prefix: str) -> None:
    """Update reference list paragraph: set content to new_prefix + rest (preserves list entry body)."""
    if _update_bookmarked_number(p, new_num):
        return
    full_text = (p.text or "").strip()
    m = _match_ref_line(full_text)
    if not m or m[0] != old_num:
        return
    _, _, rest, _ = m
    new_text = new_prefix + rest
    p.clear()
    p.add_run(new_text)


def _update_ref_field_results(p, mapping: Dict[str, str]) -> None:
    numbers = _extract_ref_numbers_in_order(p, set(mapping.keys()))
    if not numbers:
        return
    mapped = [mapping.get(n, n) for n in numbers]
    idx = 0
    in_result = False
    for r in _iter_run_elements(p):
        fld_char = r.find(qn("w:fldChar"))
        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "separate":
                in_result = True
            elif fld_type == "end":
                in_result = False
        if in_result and idx < len(mapped):
            for t in r.findall(qn("w:t")):
                if t.text and t.text.strip().isdigit():
                    t.text = mapped[idx]
                    idx += 1
                    if idx >= len(mapped):
                        break


def _replace_plain_citations(p, mapping: Dict[str, str]) -> None:
    for r in _iter_run_elements(p):
        t = r.find(qn("w:t"))
        if t is None or not t.text:
            continue
        def _sub(match):
            old = match.group(1)
            return f"[{mapping.get(old, old)}]"
        if REF_PATTERN.search(t.text):
            t.text = REF_PATTERN.sub(_sub, t.text)


def reorder_references_by_citation(
    input_docx: str,
    output_docx: str,
    ref_heading: str,
    stop_headings: Set[str],
) -> Tuple[int, int]:
    doc = Document(input_docx)
    start_idx = _find_reference_section_start(doc, ref_heading)
    if start_idx < 0:
        raise RuntimeError("Reference heading not found or no entries detected.")
    end_idx = _find_reference_section_end(doc, start_idx, stop_headings)

    ref_entries = _reference_entries_in_list(doc, start_idx, end_idx)
    ref_numbers_list = [e[0] for e in ref_entries]
    ref_numbers_set = set(ref_numbers_list)

    order = _build_first_appearance_order(doc, start_idx, ref_numbers_set)
    for num in ref_numbers_list:
        if num not in order:
            order.append(num)

    mapping = {old: str(idx + 1) for idx, old in enumerate(order)}

    ref_paragraphs: List[Tuple[str, object, str]] = []
    for p in _iter_reference_paragraphs(doc, start_idx, end_idx):
        m = _match_ref_line((p.text or "").strip())
        if not m:
            continue
        old_num, _prefix, rest, pat_idx = m
        new_num = mapping.get(old_num, old_num)
        fmt_fn = REF_LINE_PATTERNS[pat_idx][1]
        new_prefix = fmt_fn(new_num)
        _update_reference_number_text(p, old_num, new_num, new_prefix)
        ref_paragraphs.append((old_num, p._p, rest))

    for i in range(0, start_idx):
        p = doc.paragraphs[i]
        _update_ref_field_results(p, mapping)
        _replace_plain_citations(p, mapping)

    ref_heading_p = _find_reference_heading_by_index(doc, start_idx)
    if ref_heading_p is None:
        raise RuntimeError("Reference heading not found.")
    parent = ref_heading_p._p.getparent()
    for _, elem, _ in ref_paragraphs:
        parent.remove(elem)

    ref_map = {old: (elem, rest) for old, elem, rest in ref_paragraphs}
    for_num_sorted = sorted(order, key=lambda n: int(mapping[n]))
    insert_after = ref_heading_p._p
    for old_num in for_num_sorted:
        item = ref_map.get(old_num)
        if item is None:
            continue
        elem, _ = item
        insert_after.addnext(elem)
        insert_after = elem

    doc.save(output_docx)
    return len(mapping), len(order)

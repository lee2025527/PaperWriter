#!/usr/bin/env python3
"""
Single entry point for reference citation optimization (GB/T 7714 顺序编码制).

Runs: (1) reorder by first appearance + reorder list, (2) sort in-paragraph citations.
Input/output are docx paths only; no project-specific paths.
"""

import argparse
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from reorder_by_first_appearance import (
    REF_PATTERN,
    _find_reference_section_start,
    _extract_ref_numbers_in_order,
    _iter_run_elements,
    reorder_references_by_citation,
)
from sort_in_paragraph import sort_citations_in_paragraphs

DEFAULT_REF_HEADING = "参考文献"
DEFAULT_STOP_HEADINGS = {"致谢", "Acknowledgements", "Acknowledgments"}


def _detect_sequential_numbering(docx_path: str, ref_heading: str, stop_headings: set) -> bool:
    """Heuristic: document uses 顺序编码制 if we find [n] or REF ref_n in body before reference section."""
    doc = Document(docx_path)
    start_idx = _find_reference_section_start(doc, ref_heading)
    if start_idx < 0:
        return False
    ref_nums = set()
    for i in range(min(start_idx, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text or ""
        for m in REF_PATTERN.finditer(text):
            ref_nums.add(m.group(1))
        for r in _iter_run_elements(p):
            instr = r.find(qn("w:instrText"))
            if instr is not None and instr.text and "REF" in instr.text and "ref_" in instr.text:
                ref_nums.add("1")
                break
    return len(ref_nums) > 0


def run(
    input_docx: str,
    output_docx: str,
    ref_heading: str = DEFAULT_REF_HEADING,
    stop_headings: set = None,
    skip_detection: bool = False,
) -> None:
    if stop_headings is None:
        stop_headings = DEFAULT_STOP_HEADINGS
    if not skip_detection and not _detect_sequential_numbering(input_docx, ref_heading, stop_headings):
        print("未检测到顺序编码制引用（文中无 [n] 或 REF 域），可能为著者-出版年制。本技能仅支持顺序编码制，已跳过。", file=sys.stderr)
        return
    if Path(input_docx).resolve() == Path(output_docx).resolve():
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp = f.name
        try:
            reorder_references_by_citation(input_docx, tmp, ref_heading, stop_headings)
            sort_citations_in_paragraphs(tmp, output_docx, ref_heading, stop_headings)
        finally:
            Path(tmp).unlink(missing_ok=True)
    else:
        reorder_references_by_citation(input_docx, output_docx, ref_heading, stop_headings)
        sort_citations_in_paragraphs(output_docx, output_docx, ref_heading, stop_headings)
    print("文献引用格式优化完成。")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="论文 Word 文档文献引用格式优化（GB/T 7714 顺序编码制）：序号按首次出现重排，文末列表一一对应，同处引用升序。"
    )
    parser.add_argument("input_docx", help="输入论文 .docx 路径")
    parser.add_argument("output_docx", help="输出 .docx 路径（可与输入相同则覆盖）")
    parser.add_argument("--ref-heading", default=DEFAULT_REF_HEADING, help="参考文献标题，如「参考文献」或「References」")
    parser.add_argument(
        "--stop-heading",
        action="append",
        default=[],
        help="参考文献节结束标题，可多次指定；默认含 致谢、Acknowledgements 等",
    )
    parser.add_argument("--in-place", action="store_true", help="覆盖输入文件（等价于 output_docx 与 input_docx 相同）")
    parser.add_argument("--skip-detection", action="store_true", help="跳过顺序编码制检测，直接执行")
    args = parser.parse_args()

    stop = set(args.stop_heading) if args.stop_heading else DEFAULT_STOP_HEADINGS
    out = args.input_docx if args.in_place else args.output_docx
    run(args.input_docx, out, ref_heading=args.ref_heading, stop_headings=stop, skip_detection=args.skip_detection)


if __name__ == "__main__":
    main()

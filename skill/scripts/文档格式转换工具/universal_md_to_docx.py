import json
import re
import os
import sys
import argparse
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# --- Field Code Helpers for TOC ---
def add_toc_field(paragraph):
    """
    Insert a Table of Contents (TOC) field into the given paragraph.
    This creates a classic Word TOC. User needs to update fields upon opening doc.
    """
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # TOC \o "1-3" \h \z \u
    # \o "1-3": Levels 1 to 3
    # \h: Hyperlinks
    # \z: Hide page numbers in web layout
    # \u: Use outline levels
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)

# --- Configuration Loaders ---
def load_config(config_path):
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Config file not found: {config_path}")
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)

# --- Styling Functions ---
def set_style_from_config(run, font_config, style_config, override_bold=None):
    # 1. Set Latin Font
    run.font.name = font_config.get("latin", "Times New Roman")
    # 2. Set CJK Font
    cjk_font = style_config.get("font_cjk", font_config.get("cjk", "宋体"))
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cjk_font)
    # 3. Size
    run.font.size = Pt(style_config.get("font_size_pt", 12))
    # 4. Bold
    if override_bold is not None:
        run.bold = override_bold
    else:
        run.bold = style_config.get("bold", False)

def apply_paragraph_format(p, style_config):
    # Alignment
    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    align_str = style_config.get("alignment", "JUSTIFY")
    p.alignment = align_map.get(align_str, WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # Spacing
    if "space_before_pt" in style_config:
        p.paragraph_format.space_before = Pt(style_config["space_before_pt"])
    if "space_after_pt" in style_config:
        p.paragraph_format.space_after = Pt(style_config["space_after_pt"])
        
    # Line Spacing
    line_spacing = style_config.get("line_spacing", 1.5)
    if line_spacing == 1.0:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
    # Indentation
    first_indent_chars = style_config.get("first_line_indent_chars", 0)
    if first_indent_chars > 0:
        p.paragraph_format.first_line_indent = Pt(12 * first_indent_chars)
    else:
        p.paragraph_format.first_line_indent = Pt(0)
         
    hanging_chars = style_config.get("hanging_indent_chars", 0)
    if hanging_chars > 0:
        indent_val = 10.5 * hanging_chars
        p.paragraph_format.first_line_indent = Pt(-indent_val)
        p.paragraph_format.left_indent = Pt(indent_val)

# --- Main Conversion Logic ---
def convert_md_to_docx(md_path, docx_path, config_path, template_path=None):
    print(f"[Universal Converter V2] Starting...")
    print(f"Input: {md_path}")
    print(f"Config: {config_path}")
    if template_path:
        print(f"Template: {template_path}")
    else:
        print("Template: None (Creating Blank)")
    
    config = load_config(config_path)
    page_setup = config.get("page_setup", {})
    fonts = config.get("fonts", {})
    styles = config.get("styles", {})
    default_style = config.get("default_style", {})
    
    # 1. Initialize Document
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        # Move to end of document to append content
        # (Standard docx append adds to end)
    else:
        doc = Document()
        # Setup Page Layout only for new docs (Template already has it)
        section = doc.sections[0]
        if "width_cm" in page_setup:
            section.page_width = Cm(page_setup["width_cm"])
            section.page_height = Cm(page_setup["height_cm"])
        if "margin_top_cm" in page_setup:
            section.top_margin = Cm(page_setup["margin_top_cm"])
            section.bottom_margin = Cm(page_setup["margin_bottom_cm"])
            section.left_margin = Cm(page_setup["margin_left_cm"])
            section.right_margin = Cm(page_setup["margin_right_cm"])
    
    # 2. Read MD
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_references = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect [TOC]
        if '[TOC]' in line or '[toc]' in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("目录 (Table of Contents)")
            # Should style this as Heading 1 invisible or just bold body? 
            # Usually TOC title is nice to have.
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = fonts.get("cjk", "黑体") # Fallback
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fonts.get("cjk", "黑体"))
            
            p_toc = doc.add_paragraph()
            add_toc_field(p_toc)
            doc.add_page_break()
            continue

        # --- Standard Matching ---
        
        # 1. Main Title (# )
        if line.startswith('# '):
            clean_text = line.replace('# ', '')
            p = doc.add_paragraph()
            style_cfg = styles.get("title_main", default_style)
            apply_paragraph_format(p, style_cfg)
            run = p.add_run(clean_text)
            set_style_from_config(run, fonts, style_cfg)
            
        # 2. Heading 1 (## )
        elif line.startswith('## '):
            clean_text = line.replace('## ', '')
            p = doc.add_paragraph()
            style_cfg = styles.get("heading_1", default_style)
            apply_paragraph_format(p, style_cfg)
            run = p.add_run(clean_text)
            set_style_from_config(run, fonts, style_cfg)
            
            if "参考文献" in clean_text or "References" in clean_text:
                in_references = True
            else:
                in_references = False
                
        # 3. Heading 2 (### )
        elif line.startswith('### '):
            clean_text = line.replace('### ', '')
            p = doc.add_paragraph()
            style_cfg = styles.get("heading_2", default_style)
            apply_paragraph_format(p, style_cfg)
            run = p.add_run(clean_text)
            set_style_from_config(run, fonts, style_cfg)
            
        # 4. Heading 3 (#### )
        elif line.startswith('#### '):
            clean_text = line.replace('#### ', '')
            p = doc.add_paragraph()
            style_cfg = styles.get("heading_3", default_style)
            apply_paragraph_format(p, style_cfg)
            run = p.add_run(clean_text)
            set_style_from_config(run, fonts, style_cfg)

        # 5. List Items (* )
        elif line.startswith('* '):
            clean_text = line.replace('* ', '')
            p = doc.add_paragraph()
            style_cfg = styles.get("list_item", default_style)
            apply_paragraph_format(p, style_cfg)
            
            parts = clean_text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                is_bold = (i % 2 == 1)
                set_style_from_config(run, fonts, style_cfg, override_bold=is_bold)
                
        # 6. Page Break (---)
        elif line.startswith('---'):
            doc.add_page_break()
            
        # 7. References ([X])
        elif in_references and line.startswith('['):
            p = doc.add_paragraph()
            style_cfg = styles.get("references", default_style)
            apply_paragraph_format(p, style_cfg)
            run = p.add_run(line)
            set_style_from_config(run, fonts, style_cfg)
            
        # 8. Metadata Key-Value (**Key**)
        elif line.startswith('**') and '：' in line:
            p = doc.add_paragraph()
            apply_paragraph_format(p, default_style) 
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            clean_text = line.replace('**', '')
            run = p.add_run(clean_text)
            set_style_from_config(run, fonts, default_style)
            
        # 9. Normal Body Text
        else:
            p = doc.add_paragraph()
            apply_paragraph_format(p, default_style)
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                is_bold = (i % 2 == 1)
                set_style_from_config(run, fonts, default_style, override_bold=is_bold)
                
    doc.save(docx_path)
    print(f"[Success] Output saved to: {docx_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Universal Markdown to Docx Converter")
    parser.add_argument("input_md", help="Path to input Markdown file")
    parser.add_argument("output_docx", help="Path to output DOCX file")
    parser.add_argument("config_json", help="Path to styling configuration JSON")
    parser.add_argument("--template", help="Optional path to base template DOCX", default=None)
    
    args = parser.parse_args()
    
    convert_md_to_docx(args.input_md, args.output_docx, args.config_json, args.template)

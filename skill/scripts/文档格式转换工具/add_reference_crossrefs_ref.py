#!/usr/bin/env python3
"""
Add Word-native cross-references for in-text citations.

This post-processing script adds bookmarks on reference numbers and replaces
in-text [n] markers with REF fields that point to those bookmarks. It does not
modify paragraph styles or layout.
"""

import argparse
import re
from copy import deepcopy
from typing import Dict, Iterable, List, Optional, Set, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REF_PATTERN = re.compile(r"\[(\d{1,4})\]")
REF_LINE_PATTERN = re.compile(r"^\[(\d{1,4})\]")


def _iter_reference_paragraphs(doc: Document,
                               ref_heading: str,
                               stop_headings: Set[str]) -> Iterable:
    in_refs = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text == ref_heading:
            in_refs = True
            continue
        if in_refs and text in stop_headings:
            break
        if in_refs:
            yield p


def _find_reference_numbers(doc: Document,
                            ref_heading: str,
                            stop_headings: Set[str]) -> Set[str]:
    nums: Set[str] = set()
    for p in _iter_reference_paragraphs(doc, ref_heading, stop_headings):
        m = REF_LINE_PATTERN.match((p.text or "").strip())
        if m:
            nums.add(m.group(1))
    return nums


def _collect_bookmark_ids(doc: Document) -> Set[int]:
    ids: Set[int] = set()
    for bm in doc.element.xpath(".//w:bookmarkStart"):
        val = bm.get(qn("w:id"))
        if val and val.isdigit():
            ids.add(int(val))
    return ids


def _next_bookmark_id(existing: Set[int]) -> int:
    return (max(existing) + 1) if existing else 1


def _clear_paragraph_runs(p) -> None:
    for child in list(p._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            p._p.remove(child)


def _make_run(text: str, src_run) -> OxmlElement:
    run = OxmlElement("w:r")
    if src_run is not None and src_run._element.rPr is not None:
        run.append(deepcopy(src_run._element.rPr))
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def _make_field_char(char_type: str, src_run) -> OxmlElement:
    run = OxmlElement("w:r")
    if src_run is not None and src_run._element.rPr is not None:
        run.append(deepcopy(src_run._element.rPr))
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), char_type)
    run.append(fld_char)
    return run


def _make_instr_text(text: str, src_run) -> OxmlElement:
    run = OxmlElement("w:r")
    if src_run is not None and src_run._element.rPr is not None:
        run.append(deepcopy(src_run._element.rPr))
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = text
    run.append(instr)
    return run


def _build_ref_field_runs(num: str, src_run) -> List[OxmlElement]:
    runs: List[OxmlElement] = []
    runs.append(_make_run("[", src_run))
    runs.append(_make_field_char("begin", src_run))
    runs.append(_make_instr_text(f" REF ref_{num} \\h ", src_run))
    runs.append(_make_field_char("separate", src_run))
    runs.append(_make_run(num, src_run))
    runs.append(_make_field_char("end", src_run))
    runs.append(_make_run("]", src_run))
    return runs


def _add_bookmark_around_number(p, num: str, bookmark_id: int) -> None:
    text = p.text or ""
    m = REF_LINE_PATTERN.match(text.strip())
    if not m:
        return

    end_idx = text.find("]") + 1
    if end_idx <= 0:
        return

    src_run = p.runs[0] if p.runs else None
    prefix = "["
    suffix = text[end_idx:]

    run_prefix = _make_run(prefix, src_run)
    run_num = _make_run(num, src_run)
    run_suffix = _make_run("]" + suffix, src_run)

    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), f"ref_{num}")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    _clear_paragraph_runs(p)
    p._p.append(run_prefix)
    p._p.append(bookmark_start)
    p._p.append(run_num)
    p._p.append(bookmark_end)
    p._p.append(run_suffix)


def _replace_in_text_citations(doc: Document,
                               ref_heading: str,
                               anchors: Set[str]) -> int:
    linked = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text == ref_heading:
            break
        if not text or not REF_PATTERN.search(text):
            continue

        new_children: List[OxmlElement] = []
        changed = False
        for run in p.runs:
            if not run.text:
                continue
            parts = REF_PATTERN.split(run.text)
            idx = 0
            while idx < len(parts):
                if idx + 1 < len(parts) and parts[idx + 1].isdigit():
                    before = parts[idx]
                    num = parts[idx + 1]
                    if before:
                        new_children.append(_make_run(before, run))
                    if num in anchors:
                        new_children.extend(_build_ref_field_runs(num, run))
                        linked += 1
                        changed = True
                    else:
                        new_children.append(_make_run(f"[{num}]", run))
                    idx += 2
                else:
                    new_children.append(_make_run(parts[idx], run))
                    idx += 1

        if changed:
            _clear_paragraph_runs(p)
            for child in new_children:
                p._p.append(child)
    return linked


def add_crossrefs(input_docx: str,
                  output_docx: str,
                  ref_heading: str,
                  stop_headings: Set[str]) -> Tuple[int, int]:
    doc = Document(input_docx)
    anchors = _find_reference_numbers(doc, ref_heading, stop_headings)

    existing_ids = _collect_bookmark_ids(doc)
    bookmark_id = _next_bookmark_id(existing_ids)

    for p in _iter_reference_paragraphs(doc, ref_heading, stop_headings):
        m = REF_LINE_PATTERN.match((p.text or "").strip())
        if not m:
            continue
        num = m.group(1)
        _add_bookmark_around_number(p, num, bookmark_id)
        bookmark_id += 1

    linked = _replace_in_text_citations(doc, ref_heading, anchors)
    doc.save(output_docx)
    return len(anchors), linked


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Insert Word-native REF cross-references for [n] citations."
    )
    parser.add_argument("input_docx", help="Path to input DOCX file")
    parser.add_argument("output_docx", help="Path to output DOCX file")
    parser.add_argument("--ref-heading", default="参考文献", help="Reference heading text")
    parser.add_argument(
        "--stop-heading",
        action="append",
        default=["致谢", "Acknowledgements", "Acknowledgments"],
        help="Stop processing references when this heading is reached",
    )
    args = parser.parse_args()

    ref_count, link_count = add_crossrefs(
        args.input_docx,
        args.output_docx,
        args.ref_heading,
        set(args.stop_heading),
    )
    print(f"References found: {ref_count}")
    print(f"REF fields inserted: {link_count}")


if __name__ == "__main__":
    main()

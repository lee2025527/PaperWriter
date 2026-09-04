
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class AcademicTableBuilder:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_table(self, name, headers, rows):
        """生成符合学术规范的 Docx 表格"""
        doc = Document()
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid' # 后续可升级为自定义三线表样式
        
        # Header formatting
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(10.5) # 五号字
            
        # Data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if len(p.runs) > 0:
                    p.runs[0].font.size = Pt(10.5)

        output_path = os.path.join(self.output_dir, f"{name}.docx")
        doc.save(output_path)
        return output_path

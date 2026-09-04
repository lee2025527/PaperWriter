import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def build_table_docx(filename, headers, data, output_dir='output/deliver/论文图表'):
    """
    通用表格生成器。接收表头和数据，生成纯结构的 Word 三线表或标准网格表。
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'  # 默认使用标准网格，排版时可重新调整为三线表
    
    # 填充表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(10.5)
                r.bold = True
                
    # 填充数据
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(10.5)
                    
    filepath = os.path.join(output_dir, f"{filename}.docx")
    doc.save(filepath)
    print(f"✅ Generated Table: {filepath}")

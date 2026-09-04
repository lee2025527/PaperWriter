"""表格生成器：输出 Word 表格 .docx。支持从 Excel/CSV 读取数据。"""
import pandas as pd
from pathlib import Path
from typing import Dict, Any, Optional

def _load_table_data(data_root: Path, data_path: str, opts: Dict) -> tuple:
    full = (data_root / data_path) if not Path(data_path).is_absolute() else Path(data_path)
    if not full.exists():
        return [], []
    try:
        if full.suffix.lower() == ".csv":
            df = pd.read_csv(full)
        else:
            df = pd.read_excel(full, engine='openpyxl')
    except Exception as e:
        print(f"Error loading table data from {full}: {e}")
        return [], []
    
    if df.empty:
        return [], []
    
    headers = df.columns.tolist()
    rows = df.values.tolist()
    return headers, rows

def generate(item: Dict[str, Any], output_dir: Path, data_root: Path) -> Optional[Path]:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Warning: python-docx not installed. Skip table generation.")
        return None

    opts = item.get("options") or {}
    data_path = item.get("data_path") or item.get("data_source")
    
    headers, rows = [], []
    if data_path and (Path(data_path).suffix in (".xlsx", ".csv") or "/" in data_path):
        headers, rows = _load_table_data(data_root, data_path, opts)
    
    if not headers and not rows:
        headers = opts.get("headers", ["指标", "数值"])
        rows = opts.get("rows", [["示例数据", "0.0"]])

    doc = Document()
    # 根据学术规范，表题在表上
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{item.get('id', '表')}  {item.get('title', '')}")
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

    all_rows = [headers] + rows
    cols = len(all_rows[0])
    table = doc.add_table(rows=len(all_rows), cols=cols)
    table.style = "Table Grid"
    
    for i, row_data in enumerate(all_rows):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.rows[i].cells[j]
                cell.text = str(cell_text)[:500]
                # 首行加粗 (表头)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    fid = item.get("id", "表")
    title = item.get("title", "表格")
    safe_title = "".join(c for c in title if c.isalnum() or c in " _-—（）")
    out_path = output_dir / f"{fid}_{safe_title}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path

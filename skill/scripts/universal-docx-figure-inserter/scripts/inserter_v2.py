import os
import re
import json
import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from matcher import AssetMatcher

class UniversalInserter:
    """通用图表注入引擎 V2.5 (鲁棒版)"""
    def __init__(self, docx_path, assets_dir, output_path, config=None):
        self.docx_path = docx_path
        self.assets_dir = assets_dir
        self.output_path = output_path
        self.matcher = AssetMatcher(assets_dir)
        
        # 默认配置：符合学术规范
        self.config = {
            "Figure": {"MaxWidthCm": 15.0, "FontSizePt": 10.5, "Bold": True, "CaptionPos": "Below"},
            "Table": {"FontSizePt": 10.5, "Bold": True, "CaptionPos": "Above"},
            "Font": "宋体",
            "PlaceholderRegex": r'\[(?:插入)?(图|表|图表)\s*(\d+(?:[.-]\d+)?)\s*[:：]\s*([^|\]\n]+)'
        }
        if config:
            self.config.update(config)
            
        self.report = {"success": [], "failed": [], "missing": []}

    def _set_style(self, run, size_pt, bold=False):
        run.font.name = self.config["Font"]
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config["Font"])
        run.font.size = Pt(size_pt)
        run.bold = bold

    def _apply_table_style(self, table):
        """尝试应用标准表格样式，若失败则保持现状"""
        try:
            table.style = 'Table Grid'
        except:
            try:
                table.style = 'Normal Table'
            except:
                pass

    def run(self):
        doc = Document(self.docx_path)
        regex = re.compile(self.config["PlaceholderRegex"])
        
        # 1. 扫描所有待处理任务
        tasks = []
        for i, para in enumerate(doc.paragraphs):
            match = regex.search(para.text)
            if match:
                tasks.append({
                    "index": i,
                    "type": match.group(1).strip(),
                    "id": match.group(2).strip(),
                    "title": match.group(3).strip()
                })

        print(f"Found {len(tasks)} valid placeholders.")

        # 2. 倒序处理任务 (核心保护逻辑：从后往前处理，索引不会乱)
        for task in reversed(tasks):
            idx = task["index"]
            p = doc.paragraphs[idx]
            a_type = task["type"]
            a_id = task["id"]
            title = task["title"]
            
            asset_file = self.matcher.find_match(a_type, a_id, title)
            if not asset_file:
                print(f"  [!] Missing asset: {a_type} {a_id}")
                self.report["missing"].append(task)
                continue

            # 清空占位符文本，准备注入
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            is_table_asset = asset_file.suffix.lower() == '.docx'
            
            try:
                if not is_table_asset:
                    # --- 图片处理 ---
                    run = p.add_run()
                    run.add_picture(str(asset_file), width=Cm(self.config["Figure"]["MaxWidthCm"]))
                    
                    # 插入题注段落 (图下)
                    caption_p = doc.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_p.add_run(f"图 {a_id} {title}")
                    self._set_style(caption_run, self.config["Figure"]["FontSizePt"], self.config["Figure"]["Bold"])
                    
                    # 将题注移动到图片段落之后
                    p._p.addnext(caption_p._p)
                    self.report["success"].append(task)
                else:
                    # --- 表格处理 ---
                    # 题注段落 (表上)
                    caption_run = p.add_run(f"表 {a_id} {title}")
                    self._set_style(caption_run, self.config["Table"]["FontSizePt"], self.config["Table"]["Bold"])
                    p.paragraph_format.space_after = Pt(6)
                    
                    # 复制表格
                    src_doc = Document(asset_file)
                    for src_table in src_doc.tables:
                        # 在当前段落后创建一个新表
                        new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
                        self._apply_table_style(new_table)
                        
                        for r_idx, row in enumerate(src_table.rows):
                            for c_idx, cell in enumerate(row.cells):
                                target_cell = new_table.cell(r_idx, c_idx)
                                target_cell.text = cell.text
                                # 居中及字体设置
                                for cp in target_cell.paragraphs:
                                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for cr in cp.runs:
                                        self._set_style(cr, 10.5)
                        
                        # 移动表格位置到题注段落之后
                        p._p.addnext(new_table._tbl)
                    self.report["success"].append(task)
                    
            except Exception as e:
                print(f"  [X] Failed to insert {a_id}: {e}")
                task["error"] = str(e)
                self.report["failed"].append(task)

        doc.save(self.output_path)
        print(f"Done. Saved to {self.output_path}")
        return self.report

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True)
    parser.add_argument("--assets-dir", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    
    inserter = UniversalInserter(args.docx, args.assets_dir, args.output)
    inserter.run()
